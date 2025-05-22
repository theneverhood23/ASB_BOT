# ВВОДНАЯ ЧАСТЬ
import telebot
from telebot import types
import re
from docx import Document
from datetime import datetime
import os
from dotenv import load_dotenv
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
from pathlib import Path  # Добавлен импорт Path
import calendar
import logging

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),  # Вывод в консоль
        logging.FileHandler('bot.log', encoding='utf-8')  # Вывод в файл
    ]
)

load_dotenv()  # Загружает переменные из файла .env
token = os.getenv("TELEGRAM_BOT_TOKEN")  # Получает токен

# Универсальный путь к папке с шаблонами
# По умолчанию: папка templates в директории скрипта
BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = Path(os.getenv("TEMPLATE_DIR", BASE_DIR / "templates"))

bot = telebot.TeleBot(token)
bot.current_stage = None

CHECKMARK_EMOJI = "\u2705"
ONE_EMOJI = "\u0031\uFE0F\u20E3"
TWO_EMOJI = "\u0032\uFE0F\u20E3"
THREE_EMOJI = "\u0033\uFE0F\u20E3"
EXCLAM_EMOJI = "\u26A0\uFE0F"
FIRE_EMOJI = "\U0001F525"
SHIP_EMOJI = "\U0001F680"
CELEBRATE_EMOJI = "\U0001F389"
CARD_EMOJI = "\U0001F4B3"
REPEAT_EMOJI = "\U0001F503"


def ask_next_question(message):
    """Автоматически определяет следующий вопрос на основе текущего этапа"""
    if bot.current_stage == "lender_name":
        bot.send_message(message.chat.id, "Укажите ФИО займодавца:")
        
    elif bot.current_stage == "lender_bank_account":
        bot.send_message(message.chat.id, "Укажите номер банковского счета займодавца (ровно 20 цифр):")
        
    elif bot.current_stage == "lender_bank_name":
        bot.send_message(message.chat.id, "Укажите банк займодавца:")
        
    elif bot.current_stage == "lender_bic":
        bot.send_message(message.chat.id, "Укажите БИК займодавца (ровно 9 цифр):")
        
    elif bot.current_stage == "lender_corr_account":
        bot.send_message(message.chat.id, "Укажите корреспондентский счет (ровно 20 цифр):")
        
    elif bot.current_stage == "lender_tin":
        bot.send_message(message.chat.id, "Укажите ИНН займодавца (ровно 10 цифр):")
        
    elif bot.current_stage == "lender_kpp":
        bot.send_message(message.chat.id, "Укажите КПП займодавца (ровно 9 цифр):")
        
    elif bot.current_stage == "lender_address":
        bot.send_message(message.chat.id, "Укажите адрес займодавца:")
        
    elif bot.current_stage == "lender_email":
        bot.send_message(message.chat.id, "Укажите email займодавца:")
        
    elif bot.current_stage == "lender_reciever":
        bot.send_message(message.chat.id, "Укажите ФИО представителя займодавца:")
        
    elif bot.current_stage == "borrower_name":
        bot.send_message(message.chat.id, "Укажите ФИО/название заемщика:")
        
    elif bot.current_stage == "borrower_bank_account":
        bot.send_message(message.chat.id, "Укажите номер счета заемщика (20 цифр):")
        
    elif bot.current_stage == "borrower_bank_name":
        bot.send_message(message.chat.id, "Укажите банк заемщика:")
        
    elif bot.current_stage == "borrower_bic":
        bot.send_message(message.chat.id, "Укажите БИК заемщика (9 цифр):")
        
    elif bot.current_stage == "borrower_corr_account":
        bot.send_message(message.chat.id, "Укажите корр. счет заемщика (20 цифр):")
        
    elif bot.current_stage == "borrower_tin":
        bot.send_message(message.chat.id, "Укажите ИНН заемщика (10 цифр):")
        
    elif bot.current_stage == "borrower_kpp":
        bot.send_message(message.chat.id, "Укажите КПП заемщика (9 цифр):")
        
    elif bot.current_stage == "borrower_address":
        bot.send_message(message.chat.id, "Укажите адрес заемщика:")
        
    elif bot.current_stage == "borrower_email":
        bot.send_message(message.chat.id, "Укажите email заемщика:")
        
    elif bot.current_stage == "borrower_reciever":
        bot.send_message(message.chat.id, "Укажите ФИО представителя заемщика:")
        
    elif bot.current_stage == "repayment_date":
        bot.send_message(message.chat.id, "Укажите срок погашения (в днях):")
        
    elif bot.current_stage == "loan_size":
        bot.send_message(message.chat.id, "Укажите сумму займа (цифрами):")
        
    elif bot.current_stage == "repayment_time_transfer":
        bot.send_message(message.chat.id, "Укажите срок возврата (рабочих дней):")
        
    elif bot.current_stage == "share_maxsize":
        bot.send_message(message.chat.id, "Укажите максимальный размер доли (%):")
        
    elif bot.current_stage == "new_investment_discount":
        bot.send_message(message.chat.id, "Укажите размер дисконта (%):")
        
    elif bot.current_stage == "penalty_size":
        bot.send_message(message.chat.id, "Укажите размер неустойки (%):")
        
    elif bot.current_stage == "notification_date":
        bot.send_message(message.chat.id, "Укажите срок уведомления (дней):")
        
    elif bot.current_stage == "notification_date_enactment":
        bot.send_message(message.chat.id, "Укажите срок вступления изменений (дней):")
        
    else:
        bot.send_message(message.chat.id, "❌ Ошибка: неизвестный этап")

# Функция для заполнения шаблона документа
def fill_template(template_path, output_path, data):
    try:
        # Проверка существования шаблона
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")
        logging.info(f"Загрузка шаблона: {template_path}")

        doc = Document(template_path)
        normalized_data = {k: str(v) if v is not None else "" for k, v in data.items()}
        logging.info(f"Нормализованные данные для шаблона: {normalized_data}")

        def replace_placeholders(text, data):
            for key, value in data.items():
                placeholders = [
                    f"{{{{{key}}}}}",
                    f"{{{{ {key} }}}}",
                    f"{{{{{key.upper()}}}}}",
                    f"{{{{ {key.upper()} }}}}"
                ]
                for placeholder in placeholders:
                    if placeholder in text:
                        text = text.replace(placeholder, value)
            return text

        # Обработка параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraph.text = replace_placeholders(paragraph.text, normalized_data)

        # Обработка таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = replace_placeholders(cell.text, normalized_data)

        # Обработка колонтитулов
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    header.text = replace_placeholders(header.text, normalized_data)
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    footer.text = replace_placeholders(footer.text, normalized_data)

        # Сохранение документа
        doc.save(output_path)
        if not os.path.exists(output_path):
            raise FileNotFoundError(f"Файл {output_path} не был создан")
        logging.info(f"Шаблон заполнен и сохранён: {output_path}")
    except Exception as e:
        logging.error(f"Ошибка в fill_template: {str(e)}")
        raise

# ЗАЦИКЛИВАЕМ РАБОТУ БОТА
# Функция для отправки кнопки "Создать новый договор"
def send_new_contract_button(message):
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    new_contract_button = types.KeyboardButton(f"{REPEAT_EMOJI} Создать новый договор")
    keyboard.add(new_contract_button)
    bot.send_message(message.chat.id, "Для создания нового договора нажмите 'Создать новый договор'.", reply_markup=keyboard)

# Обработчик для текстового сообщения "Создать новый договор"
@bot.message_handler(func=lambda message: message.text == f"{REPEAT_EMOJI} Создать новый договор")
def handle_new_contract(message):
    # Сбрасываем текущий этап и начинаем с этапа оферты
    bot.current_stage = None
    bot.send_message(message.chat.id, "Начинаем создание нового договора!")
    handle_continue_button(message)  # Вызываем функцию для оферты

# 1-ЫЙ ЭТАП РАБОТЫ БОТА
# Handler for the introduction button aka /start
@bot.message_handler(commands=["start"])
def handle_start_button(message):
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    button = types.KeyboardButton("Продолжить")
    keyboard.add(button)

    first_welcome_message = (
        "Привет!\n"
        "Я твой помощник в создании договора конвертируемого займа."
    )

    second_welcome_message = (
        "Процесс создания договора будет состоять из 3-х этапов:\n"
        f"{ONE_EMOJI}: ознакомимся с офертой и прочими условиями\n"
        f"{TWO_EMOJI}: укажем данные для создания договора конвертируемого займа\n"
        f"{THREE_EMOJI}: оплатим и скачаем договор"
    )

    bot.send_message(message.chat.id, first_welcome_message, reply_markup=keyboard)
    bot.send_message(message.chat.id, second_welcome_message, reply_markup=keyboard)

@bot.message_handler(func=lambda message: message.text == "Продолжить")
def handle_continue_button(message):
    bot.current_stage = "oferta"
    keyboard = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    agree_oferta_button = telebot.types.KeyboardButton(f"{CHECKMARK_EMOJI} Согласен")
    keyboard.add(agree_oferta_button)
    bot.send_message(message.chat.id, "Пожалуйста, ознакомьтесь с офертой и нажмите 'Согласен', чтобы продолжить.", reply_markup=keyboard)

    document_link_oferta = "https://telegra.ph/Oferta-03-04"
    bot.send_message(message.chat.id,document_link_oferta)

@bot.message_handler(func=lambda message: message.text == f"{CHECKMARK_EMOJI} Согласен" and bot.current_stage == "oferta")
def handle_agree_oferta_button(message):
    bot.current_stage = "agreement"
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    agree_agreement_button = types.KeyboardButton(f"{CHECKMARK_EMOJI} Согласен")
    keyboard.add(agree_agreement_button)
    bot.send_message(message.chat.id, "Отлично! Теперь ознакомимся с согласием на обработку персональных данных.", reply_markup=keyboard)
    bot.send_message(message.chat.id, "Если Вы согласны, нажмите 'Согласен', чтобы продолжить.", reply_markup=keyboard)
    document_link_agreement = "https://telegra.ph/Soglasie-na-obrabotku-personalnyh-dannyh-03-05"
    bot.send_message(message.chat.id,document_link_agreement, reply_markup=keyboard)

@bot.message_handler(func=lambda message: message.text == f"{CHECKMARK_EMOJI} Согласен" and bot.current_stage == "agreement")
def handle_agree_agreement_button(message):
    bot.current_stage = "price"
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    agree_price_button = types.KeyboardButton(f"{CHECKMARK_EMOJI} Согласен")
    keyboard.add(agree_price_button)
    bot.send_message(message.chat.id, "Когда договор будет готов, его необходимо будет оплатить. Его стоимость исключительно символическая, 10 рублей.", reply_markup=keyboard)
    bot.send_message(message.chat.id, f"{EXCLAM_EMOJI} Используйте созданный договор только по назначению. Если самостоятельно внести в него изменения, мы не гарантируем, что он будет отвечать Вашим интересам.", reply_markup=keyboard)

# 2-ОЙ ЭТАП РАБОТЫ БОТА

@bot.message_handler(func=lambda message: bot.current_stage == "price" and message.text == f"{CHECKMARK_EMOJI} Согласен")
def handle_agree_price_button(message):
    # Убираем клавиатуру после согласия с условиями
    keyboard = types.ReplyKeyboardRemove()
    bot.send_message(message.chat.id,"Вы согласились со всеми условиями. Давайте приступим ко второму этапу - созданию договора!", reply_markup=keyboard)

    # Создаем клавиатуру для выбора типа договора
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    button_interest_free = types.KeyboardButton("Беспроцентный")
    button_with_interest = types.KeyboardButton("Процентный")
    keyboard.add(button_interest_free, button_with_interest)

    # Запрашиваем выбор типа договора
    bot.send_message(message.chat.id, "Какой из шаблонов договора мы будем использовать?", reply_markup=keyboard)

    # Переходим на следующий этап
    bot.current_stage = "contract_type"


@bot.message_handler(func=lambda message: bot.current_stage == "contract_type" and message.text in ["Беспроцентный", "Процентный"])
def handle_contract_type_selection(message):
    # Сохраняем выбранный тип договора
    selected_contract_type = message.text
    bot.selected_contract_type = selected_contract_type  # Сохраняем выбранный тип договора

    # Убираем клавиатуру после выбора
    keyboard = types.ReplyKeyboardRemove()

    # Отправляем сообщение с подтверждением выбранного типа договора
    confirmation_message = f"Отлично! Будем использовать {selected_contract_type.lower()} договор конвертируемого займа."
    bot.send_message(message.chat.id, confirmation_message, reply_markup=keyboard)

    if selected_contract_type == "Процентный":
        # Запрашиваем размер процентов для процентного договора
        bot.send_message(message.chat.id, "Укажите размер процента, который будет начисляться на сумму займа (необходимо указать только цифру):")
        bot.current_stage = "percentage_size"
    else:
        # Показываем календарь для выбора даты
        bot.current_stage = "conclusion_date"
        ask_date(message)  # Вызываем календарь
        
def generate_calendar(year=None, month=None):
    now = datetime.now()
    if year is None: 
        year = now.year
    if month is None: 
        month = now.month
    
    markup = types.InlineKeyboardMarkup()
    
    # Заголовок (месяц и год)
    markup.row(types.InlineKeyboardButton(
        text=f"{calendar.month_name[month]} {year}", 
        callback_data="ignore"
    ))
    
    # Дни недели
    markup.row(*[
        types.InlineKeyboardButton(day, callback_data="ignore") 
        for day in ["Пн","Вт","Ср","Чт","Пт","Сб","Вс"]
    ])
    
    # Дни месяца
    for week in calendar.monthcalendar(year, month):
        row = []
        for day in week:
            btn_text = " " if day == 0 else str(day)
            cb_data = "ignore" if day == 0 else f"day_{year}_{month}_{day}"
            row.append(types.InlineKeyboardButton(btn_text, callback_data=cb_data))
        markup.row(*row)
    
    # Кнопки навигации
    prev_month = month - 1 if month > 1 else 12
    prev_year = year - 1 if month == 1 else year
    next_month = month + 1 if month < 12 else 1
    next_year = year + 1 if month == 12 else year
    
    markup.row(
        types.InlineKeyboardButton("<", callback_data=f"change_{prev_year}_{prev_month}"),
        types.InlineKeyboardButton(">", callback_data=f"change_{next_year}_{next_month}")
    )
    
    return markup

@bot.callback_query_handler(func=lambda call: call.data.startswith(("day_", "change_")))
def handle_calendar(call):
    if call.data.startswith("day_"):
        _, year, month, day = call.data.split("_")
        date_str = f"{int(day):02d}.{int(month):02d}.{year}"
        bot.conclusion_date = date_str
        
        # Удаляем календарь
        bot.delete_message(
            chat_id=call.message.chat.id,
            message_id=call.message.message_id
        )
        
        # Проверка на "прошедшую дату"
        selected_date = datetime.strptime(date_str, "%d.%m.%Y").date()
        if selected_date < datetime.now().date():
            bot.temp_conclusion_date = date_str
            markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
            markup.add("✅ Да", "❌ Нет")
            bot.send_message(
                chat_id=call.message.chat.id,
                text=f"⚠️ Дата {date_str} уже прошла. Продолжить?",
                reply_markup=markup
            )
            bot.current_stage = "confirm_conclusion_date"
        else:
            bot.send_message(
                chat_id=call.message.chat.id,
                text=f"📅 Выбрана дата: {date_str}",
                reply_markup=types.ReplyKeyboardRemove()  # Убираем клавиатуру
            )
            # Переход к следующему шагу
            bot.current_stage = "lender_name"
            ask_next_question(call.message)
            
    elif call.data.startswith("change_"):
        _, year, month = call.data.split("_")
        markup = generate_calendar(int(year), int(month))
        bot.edit_message_reply_markup(
            chat_id=call.message.chat.id,
            message_id=call.message.message_id,
            reply_markup=markup
        )

# 5. Обработчик подтверждения для прошедшей даты
@bot.message_handler(func=lambda message: 
    bot.current_stage == "confirm_conclusion_date" 
    and message.text in ["✅ Да", "❌ Нет"])
def handle_date_confirmation(message):
    if message.text == "✅ Да":
        bot.conclusion_date = bot.temp_conclusion_date
        bot.send_message(
            chat_id=message.chat.id,
            text=f"📅 Дата сохранена: {bot.conclusion_date}",
            reply_markup=types.ReplyKeyboardRemove()
        )
        # Переход к следующему шагу
        bot.current_stage = "lender_name"
        ask_next_question(message)
    else:
        bot.current_stage = "conclusion_date"
        ask_date(message)  # Показываем календарь

@bot.message_handler(func=lambda message: bot.current_stage == "conclusion_date")
def ask_date(message):
    markup = generate_calendar()
    bot.send_message(
        chat_id=message.chat.id, 
        text="📅 Выберите дату заключения договора:", 
        reply_markup=markup
    )

@bot.message_handler(func=lambda message: bot.current_stage == "percentage_size")
def handle_percentage_size(message):
    percentage_size = message.text.strip()
    # Проверяем, что введённое значение состоит только из цифр
    if percentage_size.isdigit() and int(percentage_size) > 0:
        bot.percentage_size = percentage_size
        bot.send_message(message.chat.id, f"Размер процента: {percentage_size} %.")

        # Переходим к этапу выбора даты через календарь
        bot.current_stage = "conclusion_date"
        ask_date(message)  # Показываем календарь

    else:
        # Если введены не только цифры, отправляем сообщение об ошибке
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Число должно быть больше 0. Попробуйте снова:")

# @bot.message_handler(func=lambda message: bot.current_stage == "conclusion_date")
# def handle_conclusion_date(message):
#    # Регулярное выражение для проверки формата даты (дд.мм.гггг)
#    date_pattern = r"^(0[1-9]|[12][0-9]|3[01])\.(0[1-9]|1[0-2])\.(19|20)\d{2}$"
#
#    if re.match(date_pattern, message.text):
#        # Если формат даты корректный, сохраняем её
#        temp_date = message.text
#        input_date = datetime.strptime(temp_date, "%d.%m.%Y").date()
#        current_date = datetime.now().date() # Текущая дата
#
#        # Проверяем, что введённая дата не раньше текущей
#        if input_date < current_date:
#            # Предупреждаем пользователя, что дата уже прошла
#            bot.temp_conclusion_date = temp_date  # Сохраняем временную дату
#            bot.send_message(message.chat.id,f"Дата {temp_date} уже прошла. Вы уверены, что хотите продолжить заполнение договора с этой датой?",reply_markup=create_confirmation_keyboard())
#            bot.current_stage = "confirm_conclusion_date"
#
#        else:
#            # Если дата актуальная, сохраняем её в conclusion_date
#            bot.conclusion_date = temp_date
#            bot.send_message(message.chat.id, f"Дата заключения договора: {bot.conclusion_date}.")
#            bot.send_message(message.chat.id,"Теперь заполним информацию о займодавце.\nУкажите кто будет выступать займодавцем:")
#            bot.current_stage = "lender_name"
#
#    else:
#        # Если формат даты некорректный, просим ввести снова
#        bot.send_message(message.chat.id, "Некорректный формат даты. Пожалуйста, укажите дату в формате дд.мм.гггг:")

# Функция для создания обычной клавиатуры
def create_confirmation_keyboard():
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    yes_button = types.KeyboardButton("Да, продолжаем")
    no_button = types.KeyboardButton("Нет, давай изменим дату")
    markup.add(yes_button, no_button)
    return markup

# Обработчик выбора пользователя через обычную клавиатуру
@bot.message_handler(func=lambda message: bot.current_stage == "confirm_conclusion_date")
def handle_confirmation_keyboard(message):
    user_response = message.text.strip()
    if user_response == "Да, продолжаем":
        # Если пользователь согласен продолжить с ранее введённой датой
        bot.conclusion_date = bot.temp_conclusion_date  # Присваиваем значение временной даты основному параметру
        bot.send_message(message.chat.id, f"Дата заключения договора: {bot.conclusion_date}.")
        bot.send_message(message.chat.id,"Теперь заполним информацию о взаимодавце.\nУкажите кто будет выступать займодавцем:")
        bot.current_stage = "lender_name"

    elif user_response == "Нет, давай изменим дату":
        # Если пользователь хочет изменить дату
        bot.send_message(message.chat.id, "Пожалуйста, укажите новую дату заключения договора в формате дд.мм.гггг:")
        bot.current_stage = "conclusion_date"

# ЗАПОЛНЯЕМ ИНФОРМАЦИЮ О ВЗАИМОДАВЦЕ

@bot.message_handler(func=lambda message: bot.current_stage == "lender_name")
def handle_lender_name(message):
    # Сохраняем имя взаимодавца
    lender_name = message.text.strip()
    
    # Проверка ФИО 
    if not re.match(r'^[а-яА-ЯёЁ\s-]+$', lender_name):
        bot.send_message(message.chat.id, "❌ Некорректное ФИО. Используйте только русские буквы, пробелы и дефисы.")
        return  # Прерываем функцию, если проверка не пройдена

    # Присваиваем имя параметру {{lender_name}}
    bot.lender_name = lender_name

    # Отправляем подтверждение
    bot.send_message(message.chat.id, f"Имя займодавца: {lender_name}.")

    # Запрашиваем имя заемщика
    bot.send_message(message.chat.id, "Укажите номер его банковского счета (ровно 20 цифр):")

    # Переходим к следующему этапу
    bot.current_stage = "lender_bank_account"

@bot.message_handler(func=lambda message: bot.current_stage == "lender_bank_account")
def handle_lender_bank_account(message):
    lender_bank_account = message.text.strip()
    # Проверка, что введено ровно 20 цифр
    if lender_bank_account.isdigit() and len(lender_bank_account) == 20:
        bot.lender_bank_account = lender_bank_account
        bot.send_message(message.chat.id, f"Номер банковского счета займодавца: {lender_bank_account}.")

        bot.send_message(message.chat.id, "Укажите банк займодавца:")
        bot.current_stage = "lender_bank_name"

    else:
        # Уведомляем пользователя о неверном формате и запрашиваем повторный ввод
        bot.send_message(message.chat.id,"Неверный формат. Номер банковского счета должен содержать ровно 20 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "lender_bank_name")
def handle_lender_bank_name(message):
    lender_bank_name = message.text.strip()
    bot.lender_bank_name = lender_bank_name
    bot.send_message(message.chat.id, f"Банк займодавца: {lender_bank_name}.")

    bot.send_message(message.chat.id, "Укажите БИК (ровно 9 цифр):")
    bot.current_stage = "lender_bic"

@bot.message_handler(func=lambda message: bot.current_stage == "lender_bic")
def handle_lender_bic(message):
    lender_bic = message.text.strip()
    # Проверка, что введено ровно 9 цифр
    if lender_bic.isdigit() and len(lender_bic) == 9:
        bot.lender_bic = lender_bic
        bot.send_message(message.chat.id, f"БИК: {lender_bic}.")

        bot.send_message(message.chat.id, "Укажите корреспондентский счет (ровно 20 цифр):")
        bot.current_stage = "lender_corr_account"

    else:
        # Уведомляем пользователя о неверном формате и запрашиваем повторный ввод
        bot.send_message(message.chat.id,"Неверный формат. Номер БИК должен содержать ровно 9 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "lender_corr_account")
def handle_lender_corr_account(message):
    lender_corr_account = message.text.strip()
    if lender_corr_account.isdigit() and len(lender_corr_account) == 20:
        bot.lender_corr_account = lender_corr_account
        bot.send_message(message.chat.id, f"Корреспондентский счет: {lender_corr_account}.")

        bot.send_message(message.chat.id, "Укажите ИНН займодавца (ровно 10 цифр):")
        bot.current_stage = "lender_tin"

    else:
        bot.send_message(message.chat.id, "Неверный формат. Номер корреспондентского счета должен содержать ровно 20 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "lender_tin")
def handle_lender_tin(message):
    lender_tin = message.text.strip()
    if lender_tin.isdigit() and len(lender_tin) == 10:
        bot.lender_tin = lender_tin
        bot.send_message(message.chat.id, f"ИНН займодавца: {lender_tin}.")

        bot.send_message(message.chat.id, "Укажите КПП займодавца (ровно 9 цифр):")
        bot.current_stage = "lender_kpp"

    else:
        bot.send_message(message.chat.id,"Неверный формат. Номер ИНН должен содержать ровно 10 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "lender_kpp")
def handle_lender_kpp(message):
    lender_kpp = message.text.strip()
    if lender_kpp.isdigit() and len(lender_kpp) == 9:
        bot.lender_kpp = lender_kpp
        bot.send_message(message.chat.id, f"КПП займодавца: {lender_kpp}.")

        bot.send_message(message.chat.id, "Укажите адрес займодавца для получения уведомлений:")
        bot.current_stage = "lender_address"

    else:
        bot.send_message(message.chat.id,"Неверный формат. Номер КПП должен содержать ровно 9 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "lender_address")
def handle_lender_address(message):
    lender_address = message.text.strip()
    bot.lender_address = lender_address
    bot.send_message(message.chat.id, f"Адрес займодавца: {lender_address}.")

    bot.send_message(message.chat.id, "Укажите e-mail займодавца:")
    bot.current_stage = "lender_email"

@bot.message_handler(func=lambda message: bot.current_stage == "lender_email")
def handle_lender_email(message):
    lender_email = message.text.strip()
    # Регулярное выражение для проверки email
    email_pattern = r'^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$'
    if re.match(email_pattern, lender_email):
        # Если email корректен, сохраняем его и переходим к следующему шагу
        bot.lender_email = lender_email
        bot.send_message(message.chat.id, f"Е-mail займодавца: {lender_email}.")

        bot.send_message(message.chat.id, "Укажите фамилию и инициалы представителя займодавца, который будет получать уведомления:")
        bot.current_stage = "lender_reciever"

    else:
        # Если email некорректен, просим пользователя повторить ввод
        bot.send_message(message.chat.id, "Некорректный e-mail. Пожалуйста, укажите правильный e-mail:")

@bot.message_handler(func=lambda message: bot.current_stage == "lender_reciever")
def handle_lender_reciever(message):
    lender_reciever = message.text.strip()
    bot.lender_reciever = lender_reciever
    bot.send_message(message.chat.id, f"Фамилия и инициалы представителя: {lender_reciever}.")

    bot.send_message(message.chat.id, f"{FIRE_EMOJI} Отлично, сведения о займодавце получены.")
    bot.send_message(message.chat.id,"Теперь заполним информацию о заемщике.\nУкажите кто будет выступать заемщиком:")
    bot.current_stage = "borrower_name"

# ЗАПОЛНЯЕМ ИНФОРМАЦИЮ О ЗАЕМЩИКЕ

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_name")
def handle_borrower_name(message):
    # Сохраняем имя взаимодавца
    borrower_name = message.text.strip()
    
    # Проверка ФИО
    if not re.match(r'^[а-яА-ЯёЁ\s-]+$', borrower_name):
        bot.send_message(message.chat.id, "❌ Некорректное ФИО. Используйте только русские буквы, пробелы и дефисы.")
        return  # Прерываем функцию, если проверка не пройдена

    # Присваиваем имя параметру {{borrower_name}}
    bot.borrower_name = borrower_name

    # Отправляем подтверждение
    bot.send_message(message.chat.id, f"Имя заемщика: {borrower_name}.")

    # Запрашиваем имя заемщика
    bot.send_message(message.chat.id, "Укажите номер его банковского счета (ровно 20 цифр):")

    # Переходим к следующему этапу
    bot.current_stage = "borrower_bank_account"

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_bank_account")
def handle_borrower_bank_account(message):
    borrower_bank_account = message.text.strip()
    # Проверка, что введено ровно 20 цифр
    if borrower_bank_account.isdigit() and len(borrower_bank_account) == 20:
        bot.borrower_bank_account = borrower_bank_account
        bot.send_message(message.chat.id, f"Номер банковского счета заемщика: {borrower_bank_account}.")

        bot.send_message(message.chat.id, "Укажите банк заемщика:")
        bot.current_stage = "borrower_bank_name"

    else:
        # Уведомляем пользователя о неверном формате и запрашиваем повторный ввод
        bot.send_message(message.chat.id, "Неверный формат. Номер банковского счета должен содержать ровно 20 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_bank_name")
def handle_borrower_bank_name(message):
    borrower_bank_name = message.text.strip()
    bot.borrower_bank_name = borrower_bank_name
    bot.send_message(message.chat.id, f"Банк заемщика: {borrower_bank_name}.")

    bot.send_message(message.chat.id, "Укажите БИК (ровно 9 цифр):")
    bot.current_stage = "borrower_bic"

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_bic")
def handle_borrower_bic(message):
    borrower_bic = message.text.strip()
    if borrower_bic.isdigit() and len(borrower_bic) == 9:
        bot.borrower_bic = borrower_bic
        bot.send_message(message.chat.id, f"БИК: {borrower_bic}.")

        bot.send_message(message.chat.id, "Укажите корреспондентский счет (ровно 20 цифр):")
        bot.current_stage = "borrower_corr_account"

    else:
        bot.send_message(message.chat.id, "Неверный формат. Номер БИК должен содержать ровно 9 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_corr_account")
def handle_borrower_corr_account(message):
    borrower_corr_account = message.text.strip()
    if borrower_corr_account.isdigit() and len(borrower_corr_account) == 20:
        bot.borrower_corr_account = borrower_corr_account
        bot.send_message(message.chat.id, f"Корреспондентский счет: {borrower_corr_account}.")

        bot.send_message(message.chat.id, "Укажите ИНН заемщика (ровно 10 цифр):")
        bot.current_stage = "borrower_tin"

    else:
        bot.send_message(message.chat.id, "Неверный формат. Номер корреспондентского счета должен содержать ровно 20 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_tin")
def handle_borrower_tin(message):
    borrower_tin = message.text.strip()
    if borrower_tin.isdigit() and len(borrower_tin) == 10:
        bot.borrower_tin = borrower_tin
        bot.send_message(message.chat.id, f"ИНН заемщика: {borrower_tin}.")

        bot.send_message(message.chat.id, "Укажите КПП заемщика (ровно 9 цифр):")
        bot.current_stage = "borrower_kpp"

    else:
        bot.send_message(message.chat.id, "Неверный формат. Номер ИНН должен содержать ровно 10 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_kpp")
def handle_borrower_kpp(message):
    borrower_kpp = message.text.strip()
    if borrower_kpp.isdigit() and len(borrower_kpp) == 9:
        bot.borrower_kpp = borrower_kpp
        bot.send_message(message.chat.id, f"КПП заемщика: {borrower_kpp}.")

        bot.send_message(message.chat.id, "Укажите адрес заемщика для получения уведомлений:")
        bot.current_stage = "borrower_address"

    else:
        bot.send_message(message.chat.id, "Неверный формат. Номер КПП должен содержать ровно 9 цифр. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_address")
def handle_borrower_address(message):
    borrower_address = message.text.strip()
    bot.borrower_address = borrower_address
    bot.send_message(message.chat.id, f"Адрес заемщика: {borrower_address}.")

    bot.send_message(message.chat.id, "Укажите e-mail заемщика:")
    bot.current_stage = "borrower_email"

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_email")
def handle_borrower_email(message):
    borrower_email = message.text.strip()
    email_pattern = r'^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$'
    if re.match(email_pattern, borrower_email):
        bot.borrower_email = borrower_email
        bot.send_message(message.chat.id, f"Е-mail заемщика: {borrower_email}.")

        bot.send_message(message.chat.id, "Укажите фамилию и инициалы представителя заемщика, который будет получать уведомления:")
        bot.current_stage = "borrower_reciever"

    else:
        bot.send_message(message.chat.id, "Некорректный e-mail. Пожалуйста, укажите правильный e-mail:")

@bot.message_handler(func=lambda message: bot.current_stage == "borrower_reciever")
def handle_borrower_reciever(message):
    borrower_reciever = message.text.strip()
    bot.borrower_reciever = borrower_reciever
    bot.send_message(message.chat.id, f"Фамилия и инициалы представителя: {borrower_reciever}.")

    bot.send_message(message.chat.id, f"{SHIP_EMOJI} Отлично, осталось совсем немного!")
    bot.send_message(message.chat.id,"Укажите количество дней для погашения займа (необходимо указать только цифры):")
    bot.current_stage = "repayment_date"

# ЗАПОЛНЯЕМ ПРОЧУЮ ИНФОРМАЦИЮ

@bot.message_handler(func=lambda message: bot.current_stage == "repayment_date")
def handle_repayment_date(message):
    repayment_date = message.text.strip()
    if repayment_date.isdigit() and int(repayment_date) > 0:
        bot.repayment_date = repayment_date
        bot.send_message(message.chat.id, f"Количество дней для погашения займа: {repayment_date}.")

        bot.send_message(message.chat.id,"Укажите сумму займа (необходимо указать только цифры):")
        bot.current_stage = "loan_size"

    else:
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Введенное число должно быть больше 0. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "loan_size")
def handle_loan_size(message):
    loan_size = message.text.strip()
    if loan_size.isdigit() and int(loan_size) > 0:
        bot.loan_size = loan_size
        bot.send_message(message.chat.id, f"Сумма займа: {loan_size}.")

        bot.send_message(message.chat.id,"Укажите количество рабочих дней, в течение которых сумма займа должна быть возвращена (необходимо указать только цифры):")
        bot.current_stage = "repayment_time_transfer"

    else:
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Введенное число должно быть больше 0. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "repayment_time_transfer")
def handle_repayment_time_transfer(message):
    repayment_time_transfer = message.text.strip()
    if repayment_time_transfer.isdigit() and int(repayment_time_transfer) > 0:
        bot.repayment_time_transfer = repayment_time_transfer
        bot.send_message(message.chat.id, f"Количество рабочих дней для возврата займа: {repayment_time_transfer}.")

        bot.send_message(message.chat.id,"Укажите максимальный размер доли, которую может получить займодавец в результате конвертации:")
        bot.current_stage = "share_maxsize"

    else:
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Введенное число должно быть больше 0. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "share_maxsize")
def handle_share_maxsize(message):
    share_maxsize = message.text.strip()
    if share_maxsize.isdigit() and int(share_maxsize) > 0:
        bot.share_maxsize = share_maxsize
        bot.send_message(message.chat.id, f"Размер доли: {share_maxsize} %.")

        bot.send_message(message.chat.id,"Какой дисконт к инвестиционной оценке нового раунда будет предоставлен инвестору по договору конвертируемого займа (необходимо указать только цифры):")
        bot.current_stage = "new_investment_discount"

    else:
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Введенное число должно быть больше 0. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "new_investment_discount")
def handle_new_investment_discount(message):
    new_investment_discount = message.text.strip()
    if new_investment_discount.isdigit() and int(new_investment_discount) > 0:
        bot.new_investment_discount = new_investment_discount
        discount_float = float(new_investment_discount)
        discount_coeff = 1-(discount_float/100)
        bot.discount_coeff = discount_coeff
        bot.send_message(message.chat.id, f"Размер дисконта: {new_investment_discount} %.")

        bot.send_message(message.chat.id,"Укажите размер неустойки (необходимо указать только цифры):")
        bot.current_stage = "penalty_size"

    else:
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Введенное число должно быть больше 0. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "penalty_size")
def handle_penalty_size(message):
    penalty_size = message.text.strip()
    if penalty_size.isdigit() and int(penalty_size) > 0:
        bot.penalty_size = penalty_size
        bot.send_message(message.chat.id, f"Размер неустойки: {penalty_size} %.")

        bot.send_message(message.chat.id,"Укажите количество рабочих дней, в течение которых стороны обязуются информировать друг друга об изменениях в адресах для получения уведомлений (необходимо указать только цифры):")
        bot.current_stage = "notification_date"

    else:
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Введенное число должно быть больше 0. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "notification_date")
def handle_notification_date(message):
    notification_date = message.text.strip()
    if notification_date.isdigit() and int(notification_date) > 0:
        bot.notification_date = notification_date
        bot.send_message(message.chat.id, f"Количество рабочих дней для информирования: {notification_date}.")

        bot.send_message(message.chat.id,"Укажите количество рабочих дней, в течение которых уведомление об изменениях в адресах для получения уведомлений вступает в силу (необходимо указать только цифры):")
        bot.current_stage = "notification_date_enactment"

    else:
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Введенное число должно быть больше 0. Попробуйте снова:")

@bot.message_handler(func=lambda message: bot.current_stage == "notification_date_enactment")
def handle_notification_date_enactment(message):
    notification_date_enactment = message.text.strip()
    if notification_date_enactment.isdigit() and int(notification_date_enactment) > 0:
        bot.notification_date_enactment = notification_date_enactment
        bot.send_message(message.chat.id, f"Количество рабочих дней: {notification_date_enactment}.")

        bot.send_message(message.chat.id,f"{CELEBRATE_EMOJI} Поздравляю!\nДоговор сформирован.")

    else:
        bot.send_message(message.chat.id, "Ошибка! Необходимо указать только цифры. Введенное число должно быть больше 0. Попробуйте снова:")

    # Создаём клавиатуру с кнопкой "Оплатить"
    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    pay_contract_button = types.KeyboardButton(f"{CARD_EMOJI} Оплатить")
    keyboard.add(pay_contract_button)

    # Отправляем кнопку "Оплатить"
    bot.send_message(message.chat.id, "Оплатите, нажав на кнопку 'Оплатить'. После оплаты Вы сможете скачать договор.", reply_markup=keyboard)

    # Устанавливаем флаг для ожидания оплаты
    bot.current_stage = "await_payment"

@bot.message_handler(func=lambda message: getattr(bot, "current_stage", None) == "await_payment")
def handle_post_payment(message):
    chat_id = message.chat.id
    logging.info(f"Начало обработки оплаты, chat_id: {chat_id}, текст: {message.text}, current_stage: {getattr(bot, 'current_stage', None)}")

    # Проверка текста кнопки
    expected_text = f"{CARD_EMOJI} Оплатить"
    if message.text != expected_text:
        bot.send_message(chat_id, f"Ожидалась команда '{expected_text}', получено: {message.text}. Нажмите кнопку 'Оплатить'.")
        logging.warning(f"Неверный текст команды: {message.text}, ожидалось: {expected_text}, chat_id: {chat_id}")
        return

    required_fields = [
        "conclusion_date", "lender_name", "lender_bank_account", "lender_bank_name",
        "lender_bic", "lender_corr_account", "lender_tin", "lender_kpp",
        "lender_address", "lender_email", "lender_reciever", "borrower_name",
        "borrower_bank_account", "borrower_bank_name", "borrower_bic",
        "borrower_corr_account", "borrower_tin", "borrower_kpp", "borrower_address",
        "borrower_email", "borrower_reciever", "repayment_date", "loan_size",
        "repayment_time_transfer", "share_maxsize", "discount_coeff", "penalty_size",
        "notification_date", "notification_date_enactment"
    ]
    if getattr(bot, "selected_contract_type", None) == "Процентный":
        required_fields.append("percentage_size")

    logging.info(f"Проверка полей: {required_fields}, chat_id: {chat_id}")
    missing_fields = [field for field in required_fields if not hasattr(bot, field)]
    if missing_fields:
        bot.send_message(chat_id, f"Ошибка: не заполнены поля: {', '.join(missing_fields)}")
        logging.error(f"Отсутствуют поля: {missing_fields}, chat_id: {chat_id}")
        return

    data = {field: getattr(bot, field) for field in required_fields}
    logging.info(f"Данные для шаблона: {data}, chat_id: {chat_id}")

    # Выбор шаблона
    if getattr(bot, "selected_contract_type", None) == "Беспроцентный":
        template_path = TEMPLATE_DIR / "template_dkz_interest_free.docx"
    elif getattr(bot, "selected_contract_type", None) == "Процентный":
        template_path = TEMPLATE_DIR / "template_dkz_with_interest.docx"
    else:
        bot.send_message(chat_id, "Ошибка: тип договора не выбран.")
        logging.error(f"Тип договора не выбран, chat_id: {chat_id}")
        return

    logging.info(f"Проверка шаблона: {template_path}, существует: {template_path.exists()}, chat_id: {chat_id}")
    if not template_path.exists():
        bot.send_message(chat_id, f"Шаблон договора не найден: {template_path}. Пожалуйста, обратитесь к администратору.")
        logging.error(f"Файл шаблона не найден: {template_path}, chat_id: {chat_id}")
        return

    output_path = f"contract_{chat_id}.docx"

    # Заполнение шаблона
    try:
        logging.info(f"Запуск fill_template для {output_path}, chat_id: {chat_id}")
        fill_template(template_path, output_path, data)
        if not os.path.exists(output_path):
            raise FileNotFoundError(f"Файл {output_path} не был создан")
        logging.info(f"DOCX создан: {output_path}, существует: {os.path.exists(output_path)}, chat_id: {chat_id}")
    except Exception as e:
        bot.send_message(chat_id, f"Ошибка при создании договора: {str(e)}")
        logging.error(f"Ошибка в fill_template: {str(e)}, chat_id: {chat_id}")
        return

    # Отправка DOCX документа
    try:
        with open(output_path, "rb") as file:
            bot.send_document(chat_id, file, caption="Ваш договор в формате DOCX")
        logging.info(f"DOCX отправлен: {output_path}, chat_id: {chat_id}")
    except Exception as e:
        bot.send_message(chat_id, f"Ошибка при отправке документа: {str(e)}")
        logging.error(f"Ошибка при отправке: {str(e)}, chat_id: {chat_id}")
        return

    bot.send_message(chat_id, "Спасибо за использование нашего сервиса!", reply_markup=types.ReplyKeyboardRemove())

    # Удаление временного файла
    try:
        os.remove(output_path)
        logging.info(f"Временный файл удалён: {output_path}, chat_id: {chat_id}")
    except Exception as e:
        logging.error(f"Ошибка при удалении файла: {str(e)}, chat_id: {chat_id}")

    # Кнопка для создания нового договора
    try:
        bot.send_message(chat_id, "📄 Если хотите создать новый договор, нажмите кнопку ниже.")
        send_new_contract_button(message)
    except Exception as e:
        logging.error(f"Ошибка при отправке кнопки нового договора: {str(e)}, chat_id: {chat_id}")

    bot.current_stage = None
    logging.info(f"Этап сброшен, chat_id: {chat_id}")

# Код, который постоянно обращается к серверам Telegram
bot.polling(none_stop=True)
